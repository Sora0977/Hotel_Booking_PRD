"""Apply the thesis Word formatting rules from Rules.md to a DOCX file.

The script is intentionally conservative: by default it writes a new DOCX
instead of overwriting the input document.
"""

from __future__ import annotations

import argparse
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


THESIS_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = THESIS_ROOT / "compiled_thesis.docx"
DEFAULT_OUTPUT = THESIS_ROOT / "compiled_thesis_rules.docx"
DEFAULT_TITLE = "HỆ THỐNG ĐẶT PHÒNG KHÁCH SẠN TRỰC TUYẾN"

CHAPTER_NAMES = {
    1: "GIỚI THIỆU",
    2: "PHƯƠNG PHÁP THỰC HIỆN",
    3: "THIẾT KẾ",
    4: "THỬ NGHIỆM",
    5: "KẾT LUẬN",
}

CONTENT_HEADING_RE = re.compile(r"^\s*(?:Chương\s+)?(\d+)(?:[\.:]\s*|\.\d+)", re.IGNORECASE)
CHAPTER_HEADING_RE = re.compile(r"^\s*Chương\s+(\d+)\s*[:.\-]?\s*(.*)$", re.IGNORECASE)


@dataclass(frozen=True)
class SectionPlan:
    index: int
    kind: str
    chapter_number: int | None
    header_text: str
    starts_at_paragraph: int


def configure_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, OSError):
            pass


def clear_block(block) -> None:
    for paragraph in list(block.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def set_run_font(
    run,
    *,
    name: str = "Times New Roman",
    size: int | float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    all_caps: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    if all_caps is not None:
        run.font.all_caps = all_caps


def set_paragraph_format(
    paragraph,
    *,
    before: int | float = 6,
    after: int | float = 0,
    line_spacing: float = 1.3,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if alignment is not None:
        paragraph.alignment = alignment


def ensure_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_paragraph_style(
    style,
    *,
    size: int | float,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    all_caps: bool = False,
    before: int | float = 6,
    after: int | float = 0,
    line_spacing: float = 1.3,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    font = style.font
    font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.underline = underline
    font.all_caps = all_caps

    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if alignment is not None:
        fmt.alignment = alignment


def configure_styles(doc: Document) -> None:
    configure_paragraph_style(ensure_style(doc, "Normal"), size=13)
    configure_paragraph_style(ensure_style(doc, "Body Text"), size=13)
    configure_paragraph_style(
        ensure_style(doc, "Heading 1"),
        size=24,
        bold=True,
        all_caps=True,
        before=12,
        after=6,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    configure_paragraph_style(
        ensure_style(doc, "Heading 2"),
        size=15,
        bold=True,
        all_caps=True,
        before=12,
        after=3,
    )
    configure_paragraph_style(
        ensure_style(doc, "Heading 3"),
        size=14,
        bold=True,
        before=9,
        after=3,
    )
    configure_paragraph_style(
        ensure_style(doc, "Heading 4"),
        size=13,
        underline=True,
        before=6,
        after=0,
    )
    configure_paragraph_style(
        ensure_style(doc, "TOC Title"),
        size=18,
        bold=True,
        all_caps=True,
        before=0,
        after=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    configure_paragraph_style(
        ensure_style(doc, "Caption"),
        size=13,
        bold=True,
        italic=True,
        underline=True,
        before=3,
        after=6,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


def remove_paragraph_section_breaks(doc: Document) -> None:
    body_sect_pr = doc._element.body.sectPr
    for paragraph in doc.paragraphs:
        p_pr = paragraph._p.pPr
        if p_pr is None:
            continue
        for sect_pr in list(p_pr.findall(qn("w:sectPr"))):
            if sect_pr is not body_sect_pr:
                p_pr.remove(sect_pr)


def add_section_break_after(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for old in list(p_pr.findall(qn("w:sectPr"))):
        p_pr.remove(old)

    sect_pr = OxmlElement("w:sectPr")
    section_type = OxmlElement("w:type")
    section_type.set(qn("w:val"), "nextPage")
    sect_pr.append(section_type)
    p_pr.append(sect_pr)


def paragraph_style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def is_toc_or_front_matter(paragraph) -> bool:
    text = paragraph.text.strip().upper()
    style_name = paragraph_style_name(paragraph)
    return (
        style_name == "TOC Title"
        or "MỤC LỤC" in text
        or text in {"LỜI CẢM ƠN", "NHIỆM VỤ LUẬN VĂN", "TÓM TẮT"}
    )


def chapter_number_from_text(text: str) -> int | None:
    chapter_match = CHAPTER_HEADING_RE.match(text)
    if chapter_match:
        return int(chapter_match.group(1))
    content_match = CONTENT_HEADING_RE.match(text)
    if content_match:
        return int(content_match.group(1))
    return None


def chapter_header_text(chapter_number: int, heading_text: str) -> str:
    chapter_match = CHAPTER_HEADING_RE.match(heading_text.strip())
    if chapter_match and chapter_match.group(2).strip():
        return f"CHƯƠNG {chapter_number}: {chapter_match.group(2).strip().upper()}"
    name = CHAPTER_NAMES.get(chapter_number, "")
    return f"CHƯƠNG {chapter_number}: {name}" if name else f"CHƯƠNG {chapter_number}"


def build_section_plan(doc: Document) -> list[SectionPlan]:
    plan: list[SectionPlan] = []
    current_chapter: int | None = None
    saw_content = False

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph_style_name(paragraph)
        if not text:
            continue

        chapter_number = chapter_number_from_text(text)
        is_heading = style_name.startswith("Heading") or chapter_number is not None
        if chapter_number is None or not is_heading:
            continue

        if not saw_content and index > 0:
            plan.append(
                SectionPlan(
                    index=len(plan),
                    kind="front_matter",
                    chapter_number=None,
                    header_text="",
                    starts_at_paragraph=0,
                )
            )

        if chapter_number != current_chapter:
            plan.append(
                SectionPlan(
                    index=len(plan),
                    kind="content",
                    chapter_number=chapter_number,
                    header_text=chapter_header_text(chapter_number, text),
                    starts_at_paragraph=index,
                )
            )
            current_chapter = chapter_number
            saw_content = True

    if not plan:
        plan.append(
            SectionPlan(
                index=0,
                kind="content",
                chapter_number=1,
                header_text=chapter_header_text(1, ""),
                starts_at_paragraph=0,
            )
        )

    return plan


def rebuild_sections(doc: Document, plan: list[SectionPlan]) -> None:
    remove_paragraph_section_breaks(doc)
    break_before_indexes = [item.starts_at_paragraph for item in plan[1:]]
    for paragraph_index in sorted(set(break_before_indexes), reverse=True):
        if paragraph_index <= 0:
            continue
        add_section_break_after(doc.paragraphs[paragraph_index - 1])


def clear_header_footer_references(section) -> None:
    sect_pr = section._sectPr
    for tag in ("w:headerReference", "w:footerReference"):
        for element in list(sect_pr.findall(qn(tag))):
            sect_pr.remove(element)


def add_page_number(paragraph) -> None:
    def field_element(tag: str, attrs: dict[str, str] | None = None, text: str | None = None):
        element = OxmlElement(tag)
        for key, value in (attrs or {}).items():
            element.set(qn(key), value)
        if text is not None:
            element.text = text
        return element

    elements = (
        field_element("w:fldChar", {"w:fldCharType": "begin"}),
        field_element("w:instrText", {"xml:space": "preserve"}, " PAGE "),
        field_element("w:fldChar", {"w:fldCharType": "separate"}),
    )
    for element in elements:
        run = paragraph.add_run()
        run._r.append(element)
        set_run_font(run, size=11)

    result_run = paragraph.add_run("1")
    set_run_font(result_run, size=11)

    end_run = paragraph.add_run()
    end_run._r.append(field_element("w:fldChar", {"w:fldCharType": "end"}))
    set_run_font(end_run, size=11)


def set_page_number_start(section, start: int | None) -> None:
    sect_pr = section._sectPr
    for element in list(sect_pr.findall(qn("w:pgNumType"))):
        sect_pr.remove(element)
    if start is None:
        return

    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(start))
    sect_pr.append(pg_num_type)


def add_first_page_border(section) -> None:
    sect_pr = section._sectPr
    for element in list(sect_pr.findall(qn("w:pgBorders"))):
        sect_pr.remove(element)

    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    borders.set(qn("w:display"), "firstPage")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    sect_pr.append(borders)


def configure_sections(
    doc: Document,
    plan: list[SectionPlan],
    *,
    thesis_title: str,
    cover_border: bool,
) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.gutter = Cm(0)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    for index, section in enumerate(doc.sections):
        section_plan = plan[min(index, len(plan) - 1)]
        clear_header_footer_references(section)
        section.different_first_page_header_footer = True

        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            part.is_linked_to_previous = False
            clear_block(part)

        if section_plan.kind == "front_matter":
            set_page_number_start(section, None)
            continue

        header_paragraph = section.header.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(header_paragraph, before=0, after=0, line_spacing=1.0)
        run = header_paragraph.add_run(section_plan.header_text)
        set_run_font(run, size=11, italic=True)

        footer_paragraph = section.footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(footer_paragraph, before=0, after=0, line_spacing=1.0)
        usable_width = section.page_width - section.left_margin - section.right_margin
        footer_paragraph.paragraph_format.tab_stops.add_tab_stop(
            usable_width,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.SPACES,
        )
        title_run = footer_paragraph.add_run(thesis_title.upper())
        set_run_font(title_run, size=10, italic=True)
        footer_paragraph.add_run("\t")
        add_page_number(footer_paragraph)

        if index == first_content_section_index(plan):
            set_page_number_start(section, 1)
        else:
            set_page_number_start(section, None)

    if cover_border and doc.sections:
        add_first_page_border(doc.sections[0])


def first_content_section_index(plan: list[SectionPlan]) -> int:
    for item in plan:
        if item.kind == "content":
            return item.index
    return 0


def configure_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        style_name = paragraph_style_name(paragraph)
        if style_name == "Caption" or style_name == "Figure":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name.startswith("Heading 1"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif is_toc_or_front_matter(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_paragraph_format(paragraph, before=6, after=0, line_spacing=1.3)

        for run in paragraph.runs:
            style_size = None
            if style_name == "Heading 1":
                style_size = 24
                set_run_font(run, size=style_size, bold=True, all_caps=True)
            elif style_name == "Heading 2":
                style_size = 15
                set_run_font(run, size=style_size, bold=True, all_caps=True)
            elif style_name == "Heading 3":
                style_size = 14
                set_run_font(run, size=style_size, bold=True)
            elif style_name == "Heading 4":
                style_size = 13
                set_run_font(run, size=style_size, underline=True)
            elif style_name in {"Caption", "Figure"}:
                set_run_font(run, size=13, bold=True, italic=True, underline=True)
            else:
                set_run_font(run, size=13)


def configure_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph, before=0, after=0, line_spacing=1.15)
                    for run in paragraph.runs:
                        set_run_font(run, size=12)


def apply_rules(
    input_path: Path,
    output_path: Path,
    *,
    thesis_title: str,
    overwrite: bool,
    cover_border: bool,
) -> list[SectionPlan]:
    input_path = input_path.resolve()
    output_path = output_path.resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"Không tìm thấy file DOCX: {input_path}")
    if input_path == output_path and not overwrite:
        raise ValueError("Output trùng input. Dùng --overwrite nếu muốn ghi đè.")

    working_path = input_path
    if input_path == output_path:
        backup_path = input_path.with_suffix(".before_rules.docx")
        shutil.copy2(input_path, backup_path)

    doc = Document(working_path)
    configure_styles(doc)
    plan = build_section_plan(doc)
    rebuild_sections(doc, plan)

    # Re-open after section insertion so python-docx refreshes doc.sections.
    temp_path = output_path.with_suffix(".sections.tmp.docx")
    doc.save(temp_path)
    doc = Document(temp_path)

    configure_sections(doc, plan, thesis_title=thesis_title, cover_border=cover_border)
    configure_paragraphs(doc)
    configure_tables(doc)
    doc.save(output_path)
    temp_path.unlink(missing_ok=True)
    return plan


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Định dạng DOCX luận văn theo Rules.md, gồm layout, style, header và footer.",
    )
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT, help="File DOCX nguồn.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="File DOCX sau khi định dạng.")
    parser.add_argument(
        "--title",
        default=DEFAULT_TITLE,
        help="Tên đề tài đưa vào footer. Script sẽ tự chuyển thành chữ hoa.",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Cho phép ghi đè nếu --output trùng --input; script tạo bản .before_rules.docx trước.",
    )
    parser.add_argument(
        "--no-cover-border",
        action="store_true",
        help="Không thêm page border đơn giản cho trang đầu.",
    )
    return parser.parse_args()


def main() -> int:
    configure_stdio()
    args = parse_args()
    plan = apply_rules(
        args.input,
        args.output,
        thesis_title=args.title,
        overwrite=args.overwrite,
        cover_border=not args.no_cover_border,
    )

    print(f"Đã xuất file: {args.output.resolve()}")
    print("Section/header plan:")
    for item in plan:
        label = item.kind if item.kind == "front_matter" else item.header_text
        print(f"- section {item.index}: {label} (paragraph {item.starts_at_paragraph})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
